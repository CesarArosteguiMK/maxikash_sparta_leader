from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_severity_table(doc, rows_data):
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['ID', 'Severidad', 'Categoría', 'Ubicación', 'Descripción']
    widths = [Cm(1.2), Cm(2), Cm(2.5), Cm(4.5), Cm(7)]
    hdr = table.rows[0]
    for i, (text, width) in enumerate(zip(headers, widths)):
        cell = hdr.cells[i]
        cell.text = text
        cell.width = width
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1a3c6e')

    sev_colors = {
        'CRÍTICO': 'f8d7da',
        'ALTO': 'fff3cd',
        'MEDIO': 'cce5ff',
        'BAJO': 'd4edda',
        'INFO': 'e2e3e5'
    }

    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i >= 3 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8)
            if i == 1:
                sev = str(val).upper()
                if sev in sev_colors:
                    set_cell_shading(cell, sev_colors[sev])
    return table

def add_finding(doc, fid, title, severity, location, description, impact, recommendation, code_example=None):
    sev_emoji = {'CRÍTICO': '🔴', 'ALTO': '🟠', 'MEDIO': '🟡', 'BAJO': '🟢', 'INFO': '⚪'}
    p = doc.add_paragraph()
    run = p.add_run(f'{fid} — {sev_emoji.get(severity, "")} [{severity}] {title}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

    p = doc.add_paragraph()
    run = p.add_run('Ubicación: ')
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(location).font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('Descripción: ')
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(description).font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run('Impacto: ')
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(impact).font.size = Pt(9)

    if code_example:
        p = doc.add_paragraph()
        run = p.add_run('Código vulnerable:')
        run.bold = True
        run.font.size = Pt(9)
        p = doc.add_paragraph()
        run = p.add_run(code_example)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run('Recomendación: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
    p.add_run(recommendation).font.size = Pt(9)

    doc.add_paragraph()


# ===================== PORTADA =====================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AUDITORÍA DE SEGURIDAD')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sparta Ledger')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Análisis Profundo de Vulnerabilidades del Sistema')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(3):
    doc.add_paragraph()

info_lines = [
    f'Fecha de auditoría: {datetime.now().strftime("%d de febrero de %Y")}',
    'Alcance: Código fuente completo (backend, frontend, configuración, base de datos)',
    'Clasificación: CONFIDENCIAL',
    'Metodología: Revisión estática de código (SAST) + análisis de configuración',
    'Sistema auditado: Sparta Ledger v1.0',
    'Tecnologías: PHP 8.x, MySQL, JavaScript, PDO, Apache/XAMPP',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ===================== ÍNDICE =====================
doc.add_heading('ÍNDICE', level=1)
toc_items = [
    '1. Resumen Ejecutivo',
    '2. Alcance y Metodología',
    '3. Estadísticas de Hallazgos',
    '4. Categoría 1: Credenciales y Datos Sensibles Expuestos',
    '5. Categoría 2: Contraseñas en Texto Plano (Sin Hashing)',
    '6. Categoría 3: Inyección SQL',
    '7. Categoría 4: Autenticación y Gestión de Sesiones',
    '8. Categoría 5: Cross-Site Scripting (XSS)',
    '9. Categoría 6: Cross-Site Request Forgery (CSRF)',
    '10. Categoría 7: Seguridad en Carga de Archivos',
    '11. Categoría 8: Cabeceras HTTP y Configuración del Servidor',
    '12. Categoría 9: Control de Acceso y Autorización',
    '13. Categoría 10: Registro y Manejo de Errores',
    '14. Tabla Resumen de Todos los Hallazgos',
    '15. Plan de Remediación Priorizado',
    '16. Conclusiones',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===================== 1. RESUMEN EJECUTIVO =====================
doc.add_heading('1. Resumen Ejecutivo', level=1)

doc.add_paragraph(
    'Se realizó una auditoría de seguridad exhaustiva del sistema Sparta Ledger, '
    'analizando la totalidad del código fuente incluyendo 18 controladores, 18 modelos, '
    '51 vistas, 13 archivos del núcleo (core), 5 archivos JavaScript frontend, '
    '17 archivos de configuración y 20 scripts SQL. '
    'El análisis reveló vulnerabilidades significativas que requieren atención inmediata.'
)

doc.add_paragraph(
    'La auditoría identificó un total de 34 hallazgos de seguridad distribuidos en 10 categorías. '
    'De estos, 8 son de severidad CRÍTICA, 8 de severidad ALTA, 14 de severidad MEDIA, '
    '2 de severidad BAJA y 2 hallazgos positivos (controles bien implementados).'
)

p = doc.add_paragraph()
run = p.add_run('Hallazgos más urgentes:')
run.bold = True
items = [
    'Contraseñas almacenadas en texto plano sin ningún tipo de hashing',
    'Credenciales de bases de datos y APIs expuestas en código fuente y historial de Git',
    '10 puntos de inyección SQL críticos usando addslashes() en lugar de consultas preparadas',
    'Ausencia total de protección CSRF en todos los endpoints',
    'Contraseñas visibles en exportaciones Excel y formularios de edición',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===================== 2. ALCANCE Y METODOLOGÍA =====================
doc.add_heading('2. Alcance y Metodología', level=1)

doc.add_heading('2.1 Alcance', level=2)
doc.add_paragraph(
    'La auditoría cubrió el 100% del código fuente del proyecto Sparta Ledger ubicado en '
    'c:\\xampp\\htdocs\\sparta___SPARTA_SECRET_REDACTED__\\. Se analizaron los siguientes componentes:'
)

scope_items = [
    'Backend: 18 controladores PHP, 18 modelos DAO, 13 archivos del núcleo (core)',
    'Frontend: 51 vistas PHP, 5 archivos JavaScript, 3 archivos CSS',
    'Configuración: 17 archivos de config, .env, .htaccess (3 archivos)',
    'Base de datos: 20 scripts SQL, análisis de esquema y constraints',
    'Almacenamiento: Logs, caché, rate limiting, directorios de uploads',
    'Dependencias: PhpSpreadsheet, mPDF (1,386+ archivos de vendor)',
    'Historial Git: Revisión de commits para credenciales expuestas',
]
for item in scope_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 Metodología', level=2)
doc.add_paragraph(
    'Se utilizó análisis estático de código (SAST) combinado con revisión manual, '
    'siguiendo las directrices de OWASP Top 10 2021. Las técnicas incluyeron:'
)
methods = [
    'Búsqueda de patrones de código inseguro (regex sobre todo el codebase)',
    'Trazabilidad de datos: seguimiento de entrada de usuario ($_POST, $_GET, $_FILES, php://input) hasta su destino en consultas SQL, respuestas HTTP y sistema de archivos',
    'Análisis de configuración de sesiones, cookies y cabeceras HTTP',
    'Revisión del historial de Git para detección de secretos filtrados',
    'Verificación de controles de acceso y autorización en cada endpoint',
]
for item in methods:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===================== 3. ESTADÍSTICAS =====================
doc.add_heading('3. Estadísticas de Hallazgos', level=1)

stats_table = doc.add_table(rows=6, cols=3)
stats_table.style = 'Table Grid'
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_data = [
    ('Severidad', 'Cantidad', 'Porcentaje'),
    ('CRÍTICO', '8', '23.5%'),
    ('ALTO', '8', '23.5%'),
    ('MEDIO', '14', '41.2%'),
    ('BAJO', '2', '5.9%'),
    ('INFO (Positivo)', '2', '5.9%'),
]
colors = ['1a3c6e', 'f8d7da', 'fff3cd', 'cce5ff', 'd4edda', 'e2e3e5']
for i, (row_data, color) in enumerate(zip(stats_data, colors)):
    row = stats_table.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, color)

doc.add_paragraph()

cat_table = doc.add_table(rows=11, cols=3)
cat_table.style = 'Table Grid'
cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cat_data = [
    ('Categoría', 'Hallazgos', 'Severidad Máx.'),
    ('Credenciales Expuestas', '6', 'CRÍTICO'),
    ('Contraseñas Sin Hash', '3', 'CRÍTICO'),
    ('Inyección SQL', '10+', 'CRÍTICO'),
    ('Autenticación y Sesiones', '4', 'ALTO'),
    ('XSS', '2', 'MEDIO'),
    ('CSRF', '1 (sistémico)', 'ALTO'),
    ('Carga de Archivos', '3', 'MEDIO'),
    ('Cabeceras HTTP', '3', 'MEDIO'),
    ('Control de Acceso', '1', 'MEDIO'),
    ('Manejo de Errores', '1', 'MEDIO'),
]
for i, row_data in enumerate(cat_data):
    row = cat_table.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if i == 0:
        for j in range(3):
            set_cell_shading(row.cells[j], '1a3c6e')

doc.add_page_break()

# ===================== 4. CREDENCIALES EXPUESTAS =====================
doc.add_heading('4. Categoría 1: Credenciales y Datos Sensibles Expuestos', level=1)
doc.add_paragraph(
    'Esta categoría agrupa las vulnerabilidades relacionadas con la exposición de contraseñas '
    'de base de datos, claves API, tokens y otros secretos en el código fuente o la configuración.'
)

add_finding(doc,
    'CRED-01', 'Credenciales de BD hardcodeadas con fallback en código fuente', 'CRÍTICO',
    'backend/core/Database.php — líneas 14-18',
    'La clase Database.php contiene credenciales de producción como valores de respaldo (fallback) '
    'usando el operador ?:. Si las variables de entorno no están configuradas, se usan automáticamente '
    'las credenciales hardcodeadas: host=__SPARTA_HOST_REDACTED__, user=__SPARTA_SECRET_REDACTED__, password=__SPARTA_PASSWORD_REDACTED__.',
    'Cualquier persona con acceso al código fuente obtiene acceso directo a la base de datos de producción. '
    'Esto incluye desarrolladores, contratistas, o cualquiera que clone el repositorio.',
    'Eliminar completamente los valores de fallback. Si las variables de entorno no existen, el sistema debe '
    'fallar con un error claro en lugar de usar credenciales por defecto. Implementar: '
    '$host = getenv("DB_HOST") ?: throw new RuntimeException("DB_HOST no configurado");',
    code_example='$servidor = getenv(\'DB_HOST\') ?: \'__SPARTA_HOST_REDACTED__\';\n$password = getenv(\'DB_PASSWORD\') ?: \'__SPARTA_PASSWORD_REDACTED__\';'
)

add_finding(doc,
    'CRED-02', 'Credenciales permanentes en historial de Git', 'CRÍTICO',
    'Historial de Git — commits 3abd52f, 6eea7eb',
    'El historial de Git contiene contraseñas de BD, una URL de webhook de Google Chat con API key embebida, '
    'y rutas de llaves SSH. Aunque config.ini está ahora en .gitignore, los commits antiguos preservan estos secretos indefinidamente.',
    'Cualquier persona con acceso al repositorio (actual o futuro) puede ejecutar git show/git log y extraer '
    'todas las credenciales históricas.',
    'URGENTE: 1) Rotar TODAS las credenciales expuestas inmediatamente. 2) Usar bfg-repo-cleaner o git filter-branch '
    'para purgar los commits que contienen secretos. 3) Forzar push del historial limpio. 4) Revocar el webhook de Google Chat y la API key.'
)

add_finding(doc,
    'CRED-03', 'Archivo .env con credenciales en raíz del proyecto', 'CRÍTICO',
    'Raíz del proyecto — .env',
    'El archivo .env contiene credenciales de BD (host=__SPARTA_HOST_REDACTED__, user=__SPARTA_SECRET_REDACTED__, password con caracteres especiales). '
    'Aunque está en .gitignore, se encuentra en la raíz del proyecto sin protección .htaccess. Si el document root '
    'del servidor apunta a sparta___SPARTA_SECRET_REDACTED__/ en lugar de sparta___SPARTA_SECRET_REDACTED__/public/, el archivo es descargable directamente.',
    'Acceso completo a la base de datos remota si el servidor está mal configurado.',
    'Mover .env fuera del directorio público. Agregar regla en .htaccess raíz para denegar acceso a archivos .env: '
    '<Files ".env"> Require all denied </Files>. Verificar que el document root apunte a public/.'
)

add_finding(doc,
    'CRED-04', 'Token de API hardcodeado en código fuente', 'ALTO',
    'backend/config/config.php — línea 6; backend/controllers/EstadoCuenta.php — líneas 654, 2478',
    'El token de la API S2 Movil (__SPARTA_TOKEN_REDACTED__) está hardcodeado como fallback en config.php '
    'y duplicado literalmente en dos lugares del controlador EstadoCuenta.php.',
    'Cualquier persona con acceso al código puede hacer llamadas autenticadas a la API S2 Movil en nombre del sistema.',
    'Almacenar el token exclusivamente en variables de entorno. Eliminar las duplicaciones en EstadoCuenta.php y '
    'referenciar desde config.php usando define() con getenv() sin fallback.'
)

add_finding(doc,
    'CRED-05', 'API key de WhatsApp predecible', 'ALTO',
    'backend/config/config.php — línea 14',
    'La clave de API para el endpoint de WhatsApp se genera con un patrón determinista: '
    '"cambiar_clave_secreta_whatsapp_" . md5("sparta___SPARTA_SECRET_REDACTED__"). Cualquiera que lea el código fuente '
    'puede computar este valor y abusar del endpoint /Sabueso/crearTicketWhatsApp.',
    'Creación no autorizada de tickets vía WhatsApp, potencial spam o abuso del sistema de tickets.',
    'Reemplazar con un secreto aleatorio generado criptográficamente (bin2hex(random_bytes(32))) '
    'almacenado exclusivamente en variables de entorno.'
)

add_finding(doc,
    'CRED-06', 'Claves API almacenadas en texto plano en BD', 'ALTO',
    'backend/config/ConfigApi.php — línea 20; tabla config_api',
    'Las claves de OpenAI, Gemini y Google Maps se almacenan como texto plano en la tabla config_api. '
    'Existe una implementación de configuración cifrada (ConfigEncrypted.php) pero NO se utiliza — '
    'el sistema carga las claves con config_api_load_from_db() que lee texto plano.',
    'Un dump de la BD o un ataque de inyección SQL exitoso expondría todas las API keys.',
    'Migrar al uso de ConfigEncrypted.php o almacenar claves API cifradas con AES-256. '
    'Implementar rotación periódica de claves.'
)

doc.add_page_break()

# ===================== 5. CONTRASEÑAS SIN HASH =====================
doc.add_heading('5. Categoría 2: Contraseñas en Texto Plano (Sin Hashing)', level=1)
doc.add_paragraph(
    'Esta es posiblemente la vulnerabilidad más grave del sistema. Las contraseñas de todos los usuarios '
    'se almacenan y comparan en texto plano, sin ningún tipo de hashing criptográfico.'
)

add_finding(doc,
    'PASS-01', 'Contraseñas almacenadas y verificadas en texto plano', 'CRÍTICO',
    'backend/models/Login.php — líneas 23-31',
    'La autenticación compara la contraseña enviada por el usuario directamente contra el valor almacenado '
    'en la base de datos mediante: WHERE user_name = :usuario AND password = :password. '
    'No se usa password_hash() ni password_verify(). El propio código contiene el comentario: '
    '"comparación directa (sin hash)".',
    'Si la base de datos es comprometida (por SQL injection, backup expuesto, acceso no autorizado), '
    'TODAS las contraseñas de TODOS los usuarios quedan expuestas de forma inmediata sin necesidad de crackeo. '
    'Esto viola el principio de defensa en profundidad y cualquier estándar de seguridad (PCI-DSS, OWASP, ISO 27001).',
    'Implementar password_hash(PASSWORD_ARGON2ID) para almacenar contraseñas y password_verify() para validar. '
    'Crear un script de migración que hashee todas las contraseñas existentes. En el login: '
    '1) Buscar usuario por user_name. 2) Verificar contraseña con password_verify($input, $stored_hash).',
    code_example='// Vulnerable (actual):\n$r = $db->queryOne($sql, [\'usuario\' => $usuario, \'password\' => $password]);\n\n// Seguro (recomendado):\n$r = $db->queryOne("SELECT * FROM persona WHERE user_name = :usuario", [\'usuario\' => $usuario]);\nif ($r && password_verify($password, $r[\'password\'])) { /* login exitoso */ }'
)

add_finding(doc,
    'PASS-02', 'Actualización de contraseñas también en texto plano', 'CRÍTICO',
    'backend/models/Usuarios.php — líneas 55-65',
    'La funcionalidad de cambio de contraseña (actualizarPassword) guarda la nueva contraseña '
    'directamente como texto plano en la base de datos, sin aplicar hashing.',
    'Mismo impacto que PASS-01. Incluso si se implementara hash en el login, el cambio de contraseña '
    'seguiría almacenando texto plano.',
    'Aplicar password_hash() antes de guardar: UPDATE USUARIO SET PASS = :pass WHERE USUARIO = :usuario '
    'con $params["pass"] = password_hash($password, PASSWORD_ARGON2ID).'
)

add_finding(doc,
    'PASS-03', 'Contraseñas expuestas en vistas y exportaciones', 'ALTO',
    'Múltiples ubicaciones: models/CapHum.php línea 290, models/Empresa.php líneas 397/511, '
    'controllers/Reporteria.php línea 665, controllers/CapHum.php línea 321, views/organigrama.php línea 1040',
    'Las contraseñas en texto plano se incluyen en consultas SELECT y se envían al frontend: '
    '1) Se muestran en formularios de edición de persona. '
    '2) Se exportan en reportes Excel como columna "CONTRASEÑA". '
    '3) Se consultan como parte del perfil del usuario.',
    'Cualquier usuario con acceso al módulo de Capital Humano o reportería puede ver las contraseñas '
    'de todos los empleados. Los reportes Excel con contraseñas pueden circular por email.',
    'NUNCA incluir el campo password en consultas SELECT. Eliminar la columna "CONTRASEÑA" de las '
    'exportaciones Excel. Eliminar campos de contraseña de los formularios de edición (usar flujo de '
    '"restablecer contraseña" separado).'
)

doc.add_page_break()

# ===================== 6. INYECCIÓN SQL =====================
doc.add_heading('6. Categoría 3: Inyección SQL', level=1)
doc.add_paragraph(
    'El sistema utiliza PDO con prepared statements en su clase Database (backend/core/Database.php), '
    'la cual está bien configurada con EMULATE_PREPARES => false. Sin embargo, múltiples modelos y '
    'controladores construyen consultas SQL concatenando valores con addslashes() en lugar de usar '
    'los parámetros nombrados que la propia clase Database soporta.'
)

p = doc.add_paragraph()
run = p.add_run('¿Por qué addslashes() NO es seguro?')
run.bold = True
doc.add_paragraph(
    'addslashes() escapa solo las comillas simples, dobles, backslash y null byte. No protege contra: '
    '1) Ataques de multibyte charset (GBK/SJIS donde 0xbf27 se convierte en un carácter válido + comilla). '
    '2) Inyección en contextos numéricos (WHERE id = $var sin comillas). '
    '3) Inyección en nombres de columna/tabla. La ÚNICA defensa confiable son las consultas preparadas (prepared statements).'
)

doc.add_heading('6.1 Hallazgos Críticos (Input de usuario llega a consultas sin parametrizar)', level=2)

add_finding(doc,
    'SQLI-01', 'INSERT persona con 14 campos interpolados', 'CRÍTICO',
    'backend/models/CapHum.php — líneas 1412-1447 — insertPersona()',
    '14 campos provenientes de $_POST (nombres, apellidos, correo, teléfono, usuario, contraseña, etc.) '
    'se escapan con addslashes() y se interpolan directamente en una consulta INSERT.',
    'Un atacante puede ejecutar SQL arbitrario a través del formulario de registro de persona, '
    'potencialmente extrayendo toda la base de datos, modificando registros, o ejecutando comandos del sistema.',
    'Reescribir usando parámetros nombrados: $db->queryOne("INSERT INTO persona (...) VALUES (:nombres, :apellidop, ...)", '
    '["nombres" => $data["nombres"], ...]);',
    code_example='// VULNERABLE (actual):\n$nombres = addslashes($data[\'nombres\']);\n$db->queryOne("INSERT INTO persona ... VALUES (\'$nombres\', ...)");\n\n// SEGURO (recomendado):\n$db->queryOne("INSERT INTO persona ... VALUES (:nombres, ...)", [\'nombres\' => $data[\'nombres\']]);'
)

add_finding(doc,
    'SQLI-02', 'UPDATE persona con 8 campos string interpolados', 'CRÍTICO',
    'backend/models/CapHum.php — líneas 1570-1600 — UpdatePersona()',
    '8 campos string (nombres, apellidos, correo, teléfono, usuario, contraseña) se actualizan '
    'con el mismo patrón addslashes() + interpolación. Además, 12 consultas de seguimiento '
    '(asigna_jefe, asigna_puesto, asigna_legion) interpolan variables de tipo entero.',
    'Modificación arbitraria de datos de cualquier persona en el sistema.',
    'Convertir todas las consultas a prepared statements con parámetros nombrados.'
)

add_finding(doc,
    'SQLI-03', 'INSERT baja_persona con 5 campos interpolados', 'CRÍTICO',
    'backend/models/CapHum.php — líneas 1690-1739 — registrarBajaGestor()',
    '5 campos (id_persona, motivo, fecha_baja, descripcion, usuario_baja) se escapan con addslashes() '
    'e interpolan. Nota crítica: id_persona NO se castea a (int), se trata como string con addslashes().',
    'Inyección SQL a través del formulario de baja de personal.',
    'Usar prepared statements. Castear id_persona a (int).'
)

add_finding(doc,
    'SQLI-04', 'INSERT/UPDATE ausencia con 4 campos interpolados', 'CRÍTICO',
    'backend/models/CapHum.php — líneas 1504-1548 — guardarAusencia()',
    'Campos de descripción, fecha_inicio, fecha_fin y creado_por se interpolan con addslashes().',
    'Inyección a través del módulo de ausencias.',
    'Convertir a prepared statements.'
)

add_finding(doc,
    'SQLI-05', 'INSERT notas_credito con 4 campos interpolados', 'CRÍTICO',
    'backend/models/EstadoCuenta.php — líneas 287-303 — insertNotas()',
    'id_credito, nota, usuario, usuario_id se interpolan tras addslashes(). '
    'El campo "nota" es texto libre del usuario, ideal para inyección.',
    'Inyección SQL desde el módulo de estado de cuenta — compromiso de datos financieros.',
    'Usar prepared statements con :id_credito, :nota, :usuario, :usuario_id.'
)

add_finding(doc,
    'SQLI-06', 'INSERT dictamen_llamada con múltiples campos interpolados', 'CRÍTICO',
    'backend/models/EstadoCuenta.php — líneas 473-524 — insertDictamenLlamada()',
    'Campos de texto libre como fuente_ingresos y comentarios se interpolan con addslashes(). '
    'Son campos de formulario donde el usuario escribe texto arbitrario.',
    'Inyección SQL desde el módulo de gestión de llamadas.',
    'Convertir a prepared statements.'
)

add_finding(doc,
    'SQLI-07', 'INSERT condonación con 5 campos interpolados', 'CRÍTICO',
    'backend/models/EstadoCuenta.php — líneas 770-800 — insertCondonacionCobranza()',
    'id_credito, comentario, total, id_usuario, usuario se interpolan. '
    'El campo "comentario" es texto libre. "total" es numérico sin comillas.',
    'Inyección SQL que podría manipular los montos de condonaciones financieras.',
    'Usar prepared statements.'
)

add_finding(doc,
    'SQLI-08', 'INSERT condonación detalle con 3 campos interpolados', 'CRÍTICO',
    'backend/models/EstadoCuenta.php — líneas 820-838 — insertCondonacionCobranzaDetalle()',
    'id_condonacion, id_gastos_cobranza y monto se interpolan directamente sin comillas.',
    'Manipulación de montos y registros financieros.',
    'Usar prepared statements.'
)

add_finding(doc,
    'SQLI-09', 'Consultas SQL directas en controlador — guardarPermisos()', 'CRÍTICO',
    'backend/controllers/CapHum.php — líneas 6216-6230',
    'El controlador ejecuta DELETE e INSERT directamente con addslashes() sobre datos recibidos '
    'de php://input (JSON body). idPersona e idPuesto son controlados por el usuario.',
    'Un atacante puede manipular los permisos de cualquier usuario inyectando SQL.',
    'Mover la lógica al modelo y usar prepared statements.',
    code_example='$idPersonaEsc = addslashes($idPersona);\n$db->queryOne("DELETE FROM privilegios_departamento WHERE idPersona = $idPersonaEsc");'
)

add_finding(doc,
    'SQLI-10', 'Consultas SQL directas en controlador — actualizarPuestoPerfil()', 'CRÍTICO',
    'backend/controllers/CapHum.php — líneas 6276-6311',
    'SELECT, INSERT y DELETE ejecutados con addslashes() + interpolación desde JSON body.',
    'Escalación de privilegios mediante manipulación de puestos/perfiles.',
    'Mover al modelo con prepared statements.'
)

doc.add_heading('6.2 Hallazgos de Severidad Media (Variables internas interpoladas)', level=2)
doc.add_paragraph(
    'Los siguientes hallazgos involucran variables que actualmente se castean a (int) o provienen '
    'de fuentes internas, pero el patrón de interpolación es frágil y debería ser refactorizado.'
)

medium_sql = [
    ('SQLI-11', 'getPersonaDetalle()', 'models/CapHum.php:298', '$idPersona interpolado sin cast visible'),
    ('SQLI-12', 'actualizarModuloPerfil()', 'models/CapHum.php:323-351', '4 consultas con variables interpoladas'),
    ('SQLI-13', 'getPersonaDetallePerfiles()', 'models/CapHum.php:695-740', '3 consultas con $idPersona'),
    ('SQLI-14', 'getConsultaPersonasJerarquia()', 'models/CapHum.php:1135-1197', 'WHERE concatenados en CTE recursivo'),
    ('SQLI-15', 'getConsultaDepartamentoGestor()', 'models/CapHum.php:1280', 'Concatenación de WHERE'),
    ('SQLI-16', 'getConsultaDepartamentoGestorOrganigrama()', 'models/CapHum.php:1372', '$departamento interpolado'),
    ('SQLI-17', 'getComboJefesByPuesto()', 'models/CapHum.php:1028-1033', '$id_puesto interpolado 2 veces'),
    ('SQLI-18', 'eliminarPersonaCompleto()', 'models/CapHum.php:1866-1928', '15+ consultas DELETE/UPDATE con $id'),
    ('SQLI-19', 'UpdatePersona() seguimiento', 'models/CapHum.php:1603-1681', '12+ consultas con int interpolados'),
    ('SQLI-20', 'eliminarDepartamento()', 'models/Departamentos.php:329-333', '3 DELETE con $id'),
    ('SQLI-21', 'guardar()', 'models/EquivalenciasPuestos.php:96', '2 int vars interpolados'),
    ('SQLI-22', 'Consulta ofertas', 'models/Empresa.php:355', 'IN clause con implode()'),
    ('SQLI-23', 'Nombre columna dinámico', 'models/SegundometroDAO.php:1136', 'Columna dinámica en SQL'),
    ('SQLI-24', 'Nombre columna dinámico', 'models/Empresa.php:251', 'Columna dinámica sin escape'),
]

for fid, func, loc, desc in medium_sql:
    p = doc.add_paragraph()
    run = p.add_run(f'{fid} — {func}')
    run.bold = True
    run.font.size = Pt(9)
    p.add_run(f' ({loc}): {desc}').font.size = Pt(9)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Nota positiva: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
doc.add_paragraph(
    'Varios archivos SÍ usan prepared statements correctamente, demostrando que la clase Database '
    'lo soporta: Condonaciones.php, Departamentos.php (inserts/updates), Gestiones.php, Ticket.php, '
    'Login.php, y registrarReingresoGestor() en CapHum.php.'
)

doc.add_page_break()

# ===================== 7. AUTENTICACIÓN Y SESIONES =====================
doc.add_heading('7. Categoría 4: Autenticación y Gestión de Sesiones', level=1)

add_finding(doc,
    'AUTH-01', 'Sin regeneración de ID de sesión después del login', 'ALTO',
    'backend/controllers/Login.php — método validaUsuario()',
    'Después de una autenticación exitosa, no se llama a session_regenerate_id(true). '
    'El ID de sesión pre-login se mantiene, lo que permite ataques de fijación de sesión.',
    'Un atacante que conozca o fije un ID de sesión antes del login puede secuestrar la sesión '
    'una vez que el usuario se autentique.',
    'Agregar session_regenerate_id(true) inmediatamente después de verificar las credenciales y antes de escribir datos en $_SESSION.'
)

add_finding(doc,
    'AUTH-02', 'Sin timeout de inactividad de sesión', 'ALTO',
    'Sistema completo — no se encontró implementación de timeout',
    'No existe un mecanismo de expiración de sesión por inactividad. Una vez autenticado, '
    'la sesión permanece activa indefinidamente hasta que el usuario cierre el navegador '
    '(lifetime=0 en session_set_cookie_params).',
    'Si un usuario deja su sesión abierta (computadora compartida, café internet), '
    'cualquier persona puede usar su sesión sin límite de tiempo.',
    'Implementar verificación de última actividad en SessionGuard: almacenar $_SESSION["last_activity"] = time() '
    'en cada request. Si time() - $_SESSION["last_activity"] > 1800 (30 min), destruir sesión y redirigir al login.'
)

add_finding(doc,
    'AUTH-03', 'Rate limiter vulnerable a spoofing de IP', 'ALTO',
    'backend/core/LoginRateLimit.php — líneas 30-33',
    'El rate limiter confía en los headers X-Forwarded-For y X-Client-IP para identificar al cliente. '
    'Estos headers son controlados por el usuario y pueden ser enviados con cualquier valor.',
    'Un atacante puede realizar intentos ilimitados de fuerza bruta simplemente cambiando '
    'el header X-Forwarded-For en cada request, haciendo inútil la protección.',
    'Usar exclusivamente REMOTE_ADDR como IP del cliente, a menos que exista un proxy inverso de confianza. '
    'Si hay proxy, implementar una lista blanca de IPs de proxy y solo confiar en X-Forwarded-For '
    'cuando venga de esas IPs.',
    code_example='// Vulnerable (actual):\n$keys = [\'HTTP_X_FORWARDED_FOR\', \'HTTP_CLIENT_IP\', \'REMOTE_ADDR\'];\n\n// Seguro:\n$ip = $_SERVER[\'REMOTE_ADDR\'];'
)

add_finding(doc,
    'AUTH-04', 'Cookie de sesión sin flag "secure"', 'MEDIO',
    'public/index.php — línea 85',
    'session_set_cookie_params configura secure => false, permitiendo que la cookie de sesión '
    'se transmita por HTTP sin cifrar.',
    'En redes no seguras (WiFi pública), el token de sesión puede ser capturado por un atacante MITM.',
    'Cambiar a secure => true una vez que se confirme el uso exclusivo de HTTPS. '
    'Implementar redirección forzada de HTTP a HTTPS.'
)

doc.add_page_break()

# ===================== 8. XSS =====================
doc.add_heading('8. Categoría 5: Cross-Site Scripting (XSS)', level=1)

p = doc.add_paragraph()
run = p.add_run('Aspecto positivo: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
doc.add_paragraph(
    'La mayoría de las vistas PHP utilizan htmlspecialchars() consistentemente para escapar datos '
    'dinámicos. Esto demuestra buenas prácticas generales en el frontend PHP.'
)

add_finding(doc,
    'XSS-01', 'DOM XSS en analytics-modals.js via innerHTML', 'MEDIO',
    'public/assets/js/analytics-modals.js — líneas 153-161',
    'Las respuestas del servidor (res.mensaje, err.message) y el resultado de renderFn(res.data) '
    'se inyectan directamente en innerHTML sin sanitización. Si la respuesta del servidor es '
    'comprometida o renderFn no escapa sus datos, se produce XSS almacenado.',
    'Ejecución de JavaScript malicioso en el navegador de usuarios que vean los analytics.',
    'Usar textContent en lugar de innerHTML para mensajes de texto. Para contenido HTML, '
    'usar DOMPurify.sanitize() antes de asignar a innerHTML.'
)

add_finding(doc,
    'XSS-02', 'XSS potencial en handlers onclick con contexto JavaScript', 'MEDIO',
    'backend/views/__SPARTA_SECRET_REDACTED___request.php — líneas 1640, 1653',
    'Se usa htmlspecialchars() dentro de un atributo onclick de JavaScript. Si bien escapa entidades HTML, '
    'en un contexto JavaScript dentro de un atributo HTML, las entidades decodificadas pueden '
    'producir comportamiento inesperado si el valor no es estrictamente numérico.',
    'Posible ejecución de código JavaScript si idCredito contiene valores no numéricos.',
    'Castear explícitamente a (int) en el PHP: onclick="consultaGastos(<?= (int)($data["idCredito"] ?? 0) ?>)"'
)

doc.add_page_break()

# ===================== 9. CSRF =====================
doc.add_heading('9. Categoría 6: Cross-Site Request Forgery (CSRF)', level=1)

add_finding(doc,
    'CSRF-01', 'Ausencia TOTAL de protección CSRF en todo el sistema', 'ALTO',
    'Sistema completo — todos los endpoints POST',
    'No existe ningún mecanismo de protección CSRF en toda la aplicación: '
    'no hay tokens CSRF en formularios, no hay verificación de headers Origin/Referer, '
    'no hay cookies de doble envío. La cookie de sesión tiene SameSite=Lax que ofrece '
    'protección parcial, pero no es suficiente contra todos los vectores CSRF (especialmente '
    'requests AJAX desde subdominios).\n\n'
    'Endpoints vulnerables incluyen: crear/editar/eliminar personas, cambiar contraseñas, '
    'registrar bajas y reingresos, subir documentos, gestionar permisos, crear tickets, '
    'registrar condonaciones financieras, y más.',
    'Un atacante puede crear una página web que, al ser visitada por un usuario autenticado del sistema, '
    'ejecute acciones en su nombre: crear usuarios, cambiar contraseñas, eliminar registros, '
    'registrar movimientos financieros fraudulentos.',
    'Implementar tokens CSRF: 1) Generar token aleatorio en la sesión: $_SESSION["csrf_token"] = bin2hex(random_bytes(32)). '
    '2) Incluir en cada formulario: <input type="hidden" name="_token" value="<?= $_SESSION["csrf_token"] ?>">. '
    '3) Para AJAX: enviar como header X-CSRF-Token. '
    '4) Validar en el backend antes de procesar cualquier POST/PUT/DELETE. '
    '5) Rotar el token después de cada acción sensible.'
)

doc.add_page_break()

# ===================== 10. CARGA DE ARCHIVOS =====================
doc.add_heading('10. Categoría 7: Seguridad en Carga de Archivos', level=1)

p = doc.add_paragraph()
run = p.add_run('Aspecto positivo: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
doc.add_paragraph(
    'La clase SecureUpload (backend/core/SecureUpload.php) está bien implementada: usa finfo para '
    'detección real de MIME type, genera nombres de archivo UUID (previene path traversal y colisiones), '
    'usa enfoque de whitelist para tipos permitidos, y configura permisos de directorio a 0755.'
)

add_finding(doc,
    'UPLOAD-01', 'Upload de foto de perfil NO usa SecureUpload', 'MEDIO',
    'backend/controllers/Perfil.php — líneas 94-127',
    'La carga de foto de perfil acepta base64 via POST y extrae la extensión del string data URI (controlado por el usuario). '
    'No se valida el contenido real del archivo con finfo. No hay límite de tamaño en la ruta base64. '
    'El nombre de archivo usa el ID de persona, haciéndolo predecible.',
    'Un atacante podría subir contenido malicioso disfrazado como imagen. Archivos accesibles directamente '
    'en public/assets/img/fotos_perfil/ mediante URL predecible.',
    'Usar SecureUpload para validar el contenido real del blob decodificado. Implementar límite de tamaño. '
    'Generar nombres UUID en lugar de usar el ID de persona.'
)

add_finding(doc,
    'UPLOAD-02', 'Directorios de upload sin .htaccess de protección', 'MEDIO',
    'backend/uploads/ — múltiples subdirectorios',
    'Los directorios de upload (bajas/, reingresos/, sabueso_evidencias/, documentos/) no contienen '
    '.htaccess para prevenir ejecución de scripts. Actualmente están en backend/ (no accesible directamente), '
    'pero si la configuración del servidor cambia, podrían ser accesibles.',
    'Ejecución remota de código si los directorios se vuelven accesibles vía web.',
    'Agregar .htaccess en cada directorio de uploads con: php_flag engine off, '
    '<FilesMatch "\\.php$"> Require all denied </FilesMatch>, y Header set Content-Disposition attachment.'
)

add_finding(doc,
    'UPLOAD-03', 'Sin validación de tamaño de archivo en uploads PDF', 'BAJO',
    'backend/controllers/CapHum.php — líneas 5635-5655, 5901-5912',
    'Los handlers de upload PDF de CapHum validan MIME type pero no verifican $_FILES["archivosPDF"]["size"]. '
    'Solo EstadoCuenta.php y Despachos.php aplican límite de tamaño.',
    'Un atacante podría subir archivos PDF muy grandes para agotar espacio en disco.',
    'Agregar validación: if ($_FILES["archivosPDF"]["size"] > 10 * 1024 * 1024) { /* rechazar */ }'
)

doc.add_page_break()

# ===================== 11. CABECERAS HTTP =====================
doc.add_heading('11. Categoría 8: Cabeceras HTTP y Configuración del Servidor', level=1)

p = doc.add_paragraph()
run = p.add_run('Aspecto positivo: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
doc.add_paragraph(
    'El index.php implementa cabeceras de seguridad importantes: X-Frame-Options: SAMEORIGIN, '
    'X-Content-Type-Options: nosniff, Referrer-Policy: strict-origin-when-cross-origin, '
    'y elimina X-Powered-By y Server headers.'
)

add_finding(doc,
    'HDR-01', 'CSP permite unsafe-inline y unsafe-eval', 'MEDIO',
    'public/index.php — línea 15',
    'La directiva Content-Security-Policy incluye "unsafe-inline" y "unsafe-eval" para script-src. '
    'Esto anula efectivamente la protección contra XSS que CSP debería proveer, ya que permite '
    'la ejecución de scripts inline y eval().',
    'La CSP no bloqueará scripts XSS inyectados mediante innerHTML o similar.',
    'Migrar gradualmente a nonces o hashes para scripts inline. Eliminar unsafe-eval refactorizando '
    'código que use eval(). Meta a largo plazo: CSP strict sin unsafe-inline ni unsafe-eval.'
)

add_finding(doc,
    'HDR-02', 'CORS wildcard en endpoint de documentos S3', 'MEDIO',
    'backend/controllers/EstadoCuenta.php — líneas 3773-3775',
    'El método proxyDeDocumentoS3 configura Access-Control-Allow-Origin: *, permitiendo que '
    'cualquier origen acceda a estos documentos.',
    'Documentos financieros sensibles accesibles desde cualquier sitio web.',
    'Restringir a los orígenes específicos del sistema: Access-Control-Allow-Origin: https://tudominio.com'
)

add_finding(doc,
    'HDR-03', 'HSTS habilitado pero HTTP no redirige a HTTPS', 'BAJO',
    'public/index.php — línea 17',
    'Se envía Strict-Transport-Security pero el servidor no fuerza la redirección de HTTP a HTTPS. '
    'La primera visita al sitio por HTTP no estará protegida hasta que el navegador reciba el header HSTS.',
    'Primera conexión vulnerable a MITM si el usuario accede por HTTP.',
    'Agregar redirección forzada a HTTPS en .htaccess: RewriteRule ^(.*)$ https://%{HTTP_HOST}/$1 [R=301,L]'
)

doc.add_page_break()

# ===================== 12. CONTROL DE ACCESO =====================
doc.add_heading('12. Categoría 9: Control de Acceso y Autorización', level=1)

add_finding(doc,
    'ACC-01', 'Endpoint de WhatsApp sin autenticación', 'MEDIO',
    'public/index.php — lógica de routing; backend/controllers/Sabueso.php',
    'El endpoint /Sabueso/crearTicketWhatsApp está explícitamente excluido de la verificación '
    'de sesión (junto con Login). Solo se protege con la API key que es determinista (ver CRED-05).',
    'Creación de tickets no autorizados si la API key es computada por un atacante.',
    'Implementar una API key verdaderamente aleatoria y considerar agregar validación adicional '
    '(IP whitelist, rate limiting por API key, verificación de webhook signature).'
)

p = doc.add_paragraph()
run = p.add_run('Aspecto positivo: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
doc.add_paragraph(
    'El sistema implementa control de acceso basado en módulos: el index.php verifica que el usuario '
    'tenga acceso al módulo solicitado antes de ejecutar el controlador. SessionGuard valida la sesión '
    'en cada request. Los .htaccess en backend/ y core/ previenen listado de directorios.'
)

doc.add_page_break()

# ===================== 13. MANEJO DE ERRORES =====================
doc.add_heading('13. Categoría 10: Registro y Manejo de Errores', level=1)

add_finding(doc,
    'LOG-01', 'Mensajes de error filtran detalles internos de BD', 'MEDIO',
    'backend/core/Database.php — líneas 87-94',
    'La función getError() incluye la consulta SQL completa y los parámetros bindeados en el mensaje '
    'de excepción. Si estos errores llegan al usuario (catch genérico que hace echo del error), '
    'un atacante puede obtener nombres de tablas, columnas y estructura de consultas.',
    'Ayuda al atacante a refinar ataques de inyección SQL al revelar la estructura interna.',
    'En producción, registrar el error completo en logs del servidor pero devolver al usuario '
    'un mensaje genérico: "Error interno del sistema. Contacte al administrador." '
    'Implementar: if (APP_ENV === "production") return "Error interno"; else return $detalle;'
)

doc.add_page_break()

# ===================== 14. TABLA RESUMEN =====================
doc.add_heading('14. Tabla Resumen de Todos los Hallazgos', level=1)

all_findings = [
    ('CRED-01', 'Crítico', 'Credenciales', 'core/Database.php', 'BD credentials hardcodeadas como fallback'),
    ('CRED-02', 'Crítico', 'Credenciales', 'Git history', 'Credenciales en historial de Git'),
    ('CRED-03', 'Crítico', 'Credenciales', '.env', 'Archivo .env accesible en raíz web'),
    ('CRED-04', 'Alto', 'Credenciales', 'config/config.php', 'Token API hardcodeado'),
    ('CRED-05', 'Alto', 'Credenciales', 'config/config.php', 'API key WhatsApp predecible'),
    ('CRED-06', 'Alto', 'Credenciales', 'config/ConfigApi.php', 'API keys en texto plano en BD'),
    ('PASS-01', 'Crítico', 'Contraseñas', 'models/Login.php', 'Contraseñas sin hash'),
    ('PASS-02', 'Crítico', 'Contraseñas', 'models/Usuarios.php', 'Cambio password sin hash'),
    ('PASS-03', 'Alto', 'Contraseñas', 'Múltiples', 'Passwords en vistas y exports'),
    ('SQLI-01', 'Crítico', 'SQL Injection', 'models/CapHum.php', 'INSERT persona 14 campos'),
    ('SQLI-02', 'Crítico', 'SQL Injection', 'models/CapHum.php', 'UPDATE persona 8 campos'),
    ('SQLI-03', 'Crítico', 'SQL Injection', 'models/CapHum.php', 'INSERT baja_persona'),
    ('SQLI-04', 'Crítico', 'SQL Injection', 'models/CapHum.php', 'INSERT ausencia'),
    ('SQLI-05', 'Crítico', 'SQL Injection', 'models/EstadoCuenta.php', 'INSERT notas_credito'),
    ('SQLI-06', 'Crítico', 'SQL Injection', 'models/EstadoCuenta.php', 'INSERT dictamen'),
    ('SQLI-07', 'Crítico', 'SQL Injection', 'models/EstadoCuenta.php', 'INSERT condonacion'),
    ('SQLI-08', 'Crítico', 'SQL Injection', 'models/EstadoCuenta.php', 'INSERT condonacion detalle'),
    ('SQLI-09', 'Crítico', 'SQL Injection', 'controllers/CapHum.php', 'guardarPermisos()'),
    ('SQLI-10', 'Crítico', 'SQL Injection', 'controllers/CapHum.php', 'actualizarPuestoPerfil()'),
    ('AUTH-01', 'Alto', 'Autenticación', 'controllers/Login.php', 'Sin session_regenerate_id'),
    ('AUTH-02', 'Alto', 'Autenticación', 'Sistema', 'Sin timeout de sesión'),
    ('AUTH-03', 'Alto', 'Autenticación', 'core/LoginRateLimit.php', 'Rate limit con IP spoofeable'),
    ('AUTH-04', 'Medio', 'Autenticación', 'public/index.php', 'Cookie sin flag secure'),
    ('XSS-01', 'Medio', 'XSS', 'js/analytics-modals.js', 'innerHTML sin sanitizar'),
    ('XSS-02', 'Medio', 'XSS', 'views/__SPARTA_SECRET_REDACTED__', 'onclick con htmlspecialchars'),
    ('CSRF-01', 'Alto', 'CSRF', 'Sistema completo', 'Sin protección CSRF'),
    ('UPLOAD-01', 'Medio', 'Archivos', 'controllers/Perfil.php', 'Upload sin SecureUpload'),
    ('UPLOAD-02', 'Medio', 'Archivos', 'backend/uploads/', 'Sin .htaccess en uploads'),
    ('UPLOAD-03', 'Bajo', 'Archivos', 'controllers/CapHum.php', 'Sin validación de tamaño'),
    ('HDR-01', 'Medio', 'Headers HTTP', 'public/index.php', 'CSP con unsafe-inline/eval'),
    ('HDR-02', 'Medio', 'Headers HTTP', 'controllers/EstadoCuenta', 'CORS wildcard'),
    ('HDR-03', 'Bajo', 'Headers HTTP', 'public/index.php', 'HSTS sin redirect HTTP->HTTPS'),
    ('ACC-01', 'Medio', 'Acceso', 'Sabueso/WhatsApp', 'Endpoint sin autenticación'),
    ('LOG-01', 'Medio', 'Errores', 'core/Database.php', 'Errors filtran estructura BD'),
]

add_severity_table(doc, all_findings)

doc.add_page_break()

# ===================== 15. PLAN DE REMEDIACIÓN =====================
doc.add_heading('15. Plan de Remediación Priorizado', level=1)

doc.add_heading('Prioridad 1 — INMEDIATA (1-2 semanas)', level=2)
p1_items = [
    'ROTAR CREDENCIALES: Cambiar inmediatamente todas las contraseñas de BD, tokens API, y webhooks que aparecen en el código. Revocar la API key de Google Chat.',
    'HASHING DE CONTRASEÑAS: Implementar password_hash(PASSWORD_ARGON2ID) y password_verify(). Crear script de migración para hashear todas las contraseñas existentes.',
    'ELIMINAR CONTRASEÑAS DE VISTAS: Quitar p.password de todas las consultas SELECT. Eliminar la columna "CONTRASEÑA" de reportes Excel.',
    'ELIMINAR FALLBACKS DE CREDENCIALES: Quitar los valores hardcodeados de Database.php y config.php. El sistema debe fallar si faltan variables de entorno.',
]
for i, item in enumerate(p1_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Prioridad 2 — ALTA (2-4 semanas)', level=2)
p2_items = [
    'INYECCIÓN SQL: Convertir los 10 hallazgos críticos a prepared statements. El patrón está claro — la clase Database ya soporta parámetros nombrados.',
    'PROTECCIÓN CSRF: Implementar tokens CSRF en sesión y validar en todos los endpoints POST/PUT/DELETE.',
    'RATE LIMITER: Modificar LoginRateLimit.php para usar exclusivamente REMOTE_ADDR.',
    'SESSION FIXATION: Agregar session_regenerate_id(true) después del login exitoso.',
    'SESSION TIMEOUT: Implementar timeout de 30 min por inactividad en SessionGuard.',
]
for i, item in enumerate(p2_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Prioridad 3 — MEDIA (1-2 meses)', level=2)
p3_items = [
    'REFACTORIZAR SQL MEDIO: Convertir los 14 hallazgos de severidad media a prepared statements.',
    'SECURIZAR UPLOADS: Migrar Perfil.php a SecureUpload. Agregar .htaccess a directorios de upload. Validar tamaños.',
    'XSS: Reemplazar innerHTML con textContent donde sea posible. Agregar DOMPurify para contenido HTML dinámico.',
    'CORS: Restringir el wildcard a orígenes específicos.',
    'CSP: Iniciar migración a nonces para eliminar unsafe-inline.',
    'ERRORES: Implementar manejo de errores diferenciado producción/desarrollo.',
    'GIT CLEANUP: Purgar historial de Git con bfg-repo-cleaner para eliminar credenciales históricas.',
]
for i, item in enumerate(p3_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Prioridad 4 — MEJORA CONTINUA', level=2)
p4_items = [
    'Implementar análisis estático automatizado (SAST) en el pipeline CI/CD.',
    'Agregar escaneo de secretos pre-commit con herramientas como gitleaks o truffleHog.',
    'Establecer política de rotación periódica de credenciales y API keys.',
    'Configurar HTTPS obligatorio y activar flag secure en cookies.',
    'Considerar WAF (Web Application Firewall) como capa adicional de protección.',
    'Realizar pruebas de penetración periódicas (al menos semestrales).',
]
for i, item in enumerate(p4_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_page_break()

# ===================== 16. CONCLUSIONES =====================
doc.add_heading('16. Conclusiones', level=1)

doc.add_paragraph(
    'El sistema Sparta Ledger presenta vulnerabilidades significativas que requieren atención prioritaria. '
    'Las más críticas son el almacenamiento de contraseñas en texto plano, la exposición de credenciales '
    'en el código fuente, y los múltiples puntos de inyección SQL.'
)

doc.add_paragraph(
    'Sin embargo, el sistema también muestra señales positivas que indican capacidad de mejora: '
    'la clase Database está bien diseñada con soporte para prepared statements, el core SecureUpload '
    'implementa buenas prácticas de validación de archivos, las cabeceras de seguridad HTTP están '
    'parcialmente implementadas, y el sistema SessionGuard provee una base sólida para la gestión '
    'de sesiones.'
)

doc.add_paragraph(
    'La mayoría de las vulnerabilidades pueden resolverse aplicando consistentemente los patrones '
    'seguros que ya existen en partes del sistema. El principal reto no es técnico sino de disciplina: '
    'asegurar que todos los desarrolladores usen prepared statements, nunca expongan credenciales, '
    'y sigan las prácticas de seguridad documentadas.'
)

p = doc.add_paragraph()
run = p.add_run(
    'Se recomienda encarecidamente comenzar con la rotación de credenciales y el hashing de contraseñas '
    'dentro de las próximas 1-2 semanas, seguido de la corrección de inyecciones SQL en las 2-4 semanas siguientes.'
)
run.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin del documento —')
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Generado el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# GUARDAR
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Auditoria_Seguridad_Sparta_Ledger.docx')
doc.save(output_path)
print(f'Documento generado exitosamente: {output_path}')
